#!/usr/bin/env python3
"""Resume ATS Engine — build_resume.py

Gera .docx (python-docx) e .pdf (HTML + WeasyPrint) a partir de um JSON de
currículo, com verificação automática de ATS-parseabilidade (pypdf).
O currículo é SEMPRE um documento formal (single column, sem cards/cores).

Uso:
    build_resume.py <input.json> <output_base> [--letter]

Output:
    <output_base>.docx
    <output_base>.pdf
    (verifica texto extraível + nome do candidato no PDF + nº de páginas)

JSON schema (todas as chaves opcionais exceto meta.name):
{
  "meta": {"name","email","phone","location","linkedin","github","website",
           "target_title","lang"},
  "summary": "...",
  "skills_groups": [{"title": "...", "items": ["..."]}],
  "experience": [{"title","company","location","dates","bullets":[...]}],
  "projects": [{"name","description","link"}],
  "education": [{"degree","school","dates"}],
  "certifications": ["..."]
}
"""
import json
import sys
import html as html_mod
from pathlib import Path

# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def build_docx(data, path, letter=False):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(0.75)
    if letter:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    else:
        sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)  # A4

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = Pt(12)

    usable = (sec.page_width.inches - 1.5)

    def para(text="", bold=False, size=10.5, align=None, space_after=2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if align is not None:
            p.alignment = align
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        return p, r

    def section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(11.5)
        # bottom border (sutil, ATS-safe — não é tabela)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_contact_line(meta):
        parts = []
        if meta.get("phone"): parts.append(meta["phone"])
        if meta.get("location"): parts.append(meta["location"])
        if meta.get("linkedin"): parts.append("linkedin.com/in/" + meta["linkedin"].replace("linkedin.com/in/", ""))
        if meta.get("github"): parts.append("github.com/" + meta["github"].replace("github.com/", ""))
        if meta.get("website"): parts.append(meta["website"].replace("https://", "").replace("http://", ""))
        return "  |  ".join(parts)

    def add_hyperlink(paragraph, url, text, size=10.5, bold=False):
        """Insere hyperlink clicável em parágrafo do DOCX (python-docx não tem API nativa)."""
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz)
        if bold:
            b = OxmlElement("w:b")
            rPr.append(b)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    meta = data.get("meta", {})
    # Nome
    p, r = para(meta.get("name", ""), bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    # Target title (opcional, abaixo do nome)
    if meta.get("target_title"):
        para(meta["target_title"], size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    # Contato
    contato = add_contact_line(meta)
    if contato:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        if meta.get("website"):
            r = p.add_run(contato.rsplit(" | ", 1)[0] + "  |  ")
            r.font.size = Pt(10)
            add_hyperlink(p, "https://" + meta["website"].replace("https://", "").replace("http://", ""),
                          meta["website"].replace("https://", "").replace("http://", ""), size=10)
        else:
            r = p.add_run(contato)
            r.font.size = Pt(10)

    # Summary
    if data.get("summary"):
        section_heading("Summary" if meta.get("lang", "en").lower() != "pt" else "Resumo")
        para(data["summary"].strip(), space_after=4)

    # Skills (topo para tech)
    if data.get("skills_groups"):
        section_heading("Skills")
        for grp in data["skills_groups"]:
            title = grp.get("title", "").strip()
            items = grp.get("items", [])
            if title and items:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                r1 = p.add_run(title + ": ")
                r1.bold = True
                r1.font.size = Pt(10.5)
                r2 = p.add_run(", ".join(items))
                r2.font.size = Pt(10.5)

    # Experience
    if data.get("experience"):
        section_heading("Experience" if meta.get("lang", "en").lower() != "pt" else "Experiência")
        for job in data["experience"]:
            head = " · ".join(x for x in [job.get("title"), job.get("company"), job.get("location")] if x)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(usable), WD_TAB_ALIGNMENT.RIGHT)
            r = p.add_run(head)
            r.bold = True
            r.font.size = Pt(10.5)
            if job.get("dates"):
                r2 = p.add_run("\t" + job["dates"])
                r2.font.size = Pt(10.5)
            for b in job.get("bullets", []):
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Inches(0.18)
                bp.paragraph_format.space_after = Pt(1)
                if isinstance(b, dict):
                    br = bp.add_run("• " + b.get("text", ""))
                    br.font.size = Pt(10.5)
                    if b.get("link"):
                        bp.paragraph_format.space_after = Pt(0)
                        lp = doc.add_paragraph()
                        lp.paragraph_format.left_indent = Inches(0.34)
                        lp.paragraph_format.space_after = Pt(1)
                        add_hyperlink(lp, b["link"], "→ " + b["link"].replace("https://www.instagram.com/", "instagram.com/"), size=8.5)
                else:
                    br = bp.add_run("• " + b)
                    br.font.size = Pt(10.5)

    # Projects
    if data.get("projects"):
        section_heading("Projects" if meta.get("lang", "en").lower() != "pt" else "Projetos")
        for prj in data["projects"]:
            name = prj.get("name", "")
            link = prj.get("link", "")
            desc = prj.get("description", "")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(name + ((" — " + link.replace("https://", "")) if link else ""))
            r.bold = True
            r.font.size = Pt(10.5)
            if desc:
                dp = doc.add_paragraph()
                dp.paragraph_format.left_indent = Inches(0.18)
                dp.paragraph_format.space_after = Pt(1)
                dr = dp.add_run(desc)
                dr.font.size = Pt(10.5)

    # Education
    if data.get("education"):
        section_heading("Education" if meta.get("lang", "en").lower() != "pt" else "Formação")
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(usable), WD_TAB_ALIGNMENT.RIGHT)
            head = " · ".join(x for x in [edu.get("degree"), edu.get("school")] if x)
            r = p.add_run(head)
            r.bold = True
            r.font.size = Pt(10.5)
            if edu.get("dates"):
                r2 = p.add_run("\t" + edu["dates"])
                r2.font.size = Pt(10.5)

    # Certifications
    if data.get("certifications"):
        section_heading("Certifications" if meta.get("lang", "en").lower() != "pt" else "Certificações")
        para(", ".join(data["certifications"]), space_after=2)

    doc.save(path)
    return str(path)

# ---------------------------------------------------------------------------
# PDF via HTML + WeasyPrint (mesmo conteúdo do DOCX — formal, single column)
# ---------------------------------------------------------------------------
def build_pdf(data, path, letter=False):
    import weasyprint

    meta = data.get("meta", {})
    lang_pt = meta.get("lang", "en").lower() == "pt"
    S = lambda en, pt: pt if lang_pt else en

    def esc(s):
        return html_mod.escape(str(s or ""), quote=False)

    def contact_line(meta):
        parts = []
        if meta.get("phone"): parts.append(esc(meta["phone"]))
        if meta.get("location"): parts.append(esc(meta["location"]))
        if meta.get("linkedin"): parts.append("linkedin.com/in/" + esc(meta["linkedin"].replace("linkedin.com/in/", "")))
        if meta.get("github"): parts.append("github.com/" + esc(meta["github"].replace("github.com/", "")))
        if meta.get("website"):
            url = meta["website"].replace("https://", "").replace("http://", "")
            parts.append(f'<a href="https://{url}">{esc(url)}</a>')
        return "  |  ".join(parts)

    def section(title):
        return f'<h2>{esc(title)}</h2>'

    chunks = []
    chunks.append("<!DOCTYPE html><html><head><meta charset='utf-8'><style>")
    chunks.append("""
@page { size: A4; margin: 0.75in; }
* { box-sizing: border-box; }
body { font-family: 'DejaVu Sans', sans-serif; font-size: 10.5pt; color: #000; margin: 0; }
h1 { font-size: 16pt; text-align: center; margin: 0 0 1pt; }
.role-title { font-size: 11pt; text-align: center; margin: 0 0 2pt; }
.contact { text-align: center; font-size: 10pt; margin: 0 0 8pt; }
h2 { font-size: 11.5pt; text-transform: uppercase; border-bottom: 1px solid #000;
     margin: 8pt 0 4pt; padding-bottom: 1pt; }
p.summary { margin: 0 0 6pt; }
.skill-group { margin: 0 0 2pt; }
.skill-group b { font-weight: bold; }
.job { margin: 0 0 6pt; }
.job .head { font-weight: bold; }
.job .dates { float: right; font-weight: normal; }
ul { margin: 1pt 0 4pt; padding-left: 14pt; }
li { margin: 0 0 1pt; }
.project { margin: 0 0 4pt; }
.project .pname { font-weight: bold; }
.edu { margin: 0 0 3pt; }
.edu .head { font-weight: bold; }
.edu .dates { float: right; font-weight: normal; }
.certs { margin: 0 0 4pt; }
""")
    chunks.append("</style></head><body>")
    chunks.append(f"<h1>{esc(meta.get('name',''))}</h1>")
    if meta.get("target_title"):
        chunks.append(f"<div class='role-title'>{esc(meta['target_title'])}</div>")
    c = contact_line(meta)
    if c:
        chunks.append(f"<div class='contact'>{c}</div>")

    if data.get("summary"):
        chunks.append(section(S("Summary", "Resumo")))
        chunks.append(f"<p class='summary'>{esc(data['summary'].strip())}</p>")

    if data.get("skills_groups"):
        chunks.append(section("Skills"))
        for grp in data["skills_groups"]:
            if grp.get("title") and grp.get("items"):
                chunks.append(f"<p class='skill-group'><b>{esc(grp['title'])}:</b> {esc(', '.join(grp['items']))}</p>")

    if data.get("experience"):
        chunks.append(section(S("Experience", "Experiência")))
        for job in data["experience"]:
            head = " · ".join(x for x in [job.get("title"), job.get("company"), job.get("location")] if x)
            dates = f"<span class='dates'>{esc(job.get('dates',''))}</span>" if job.get("dates") else ""
            chunks.append(f"<div class='job'><div class='head'>{esc(head)}{dates}</div>")
            if job.get("bullets"):
                chunks.append("<ul>")
                for b in job["bullets"]:
                    if isinstance(b, dict):
                        chunks.append(f"<li>{esc(b.get('text',''))}")
                        if b.get("link"):
                            url = b["link"].replace("https://www.instagram.com/", "instagram.com/")
                            chunks.append(f" <a href='{b['link']}'>(ver reel: {esc(url)})</a>")
                        chunks.append("</li>")
                    else:
                        chunks.append(f"<li>{esc(b)}</li>")
                chunks.append("</ul>")
            chunks.append("</div>")

    if data.get("projects"):
        chunks.append(section(S("Projects", "Projetos")))
        for prj in data["projects"]:
            link = prj.get("link", "")
            name = prj.get("name", "") + (f" — {link.replace('https://','').replace('http://','')}" if link else "")
            chunks.append(f"<div class='project'><span class='pname'>{esc(name)}</span>")
            if prj.get("description"):
                chunks.append(f"<div>{esc(prj['description'])}</div>")
            chunks.append("</div>")

    if data.get("education"):
        chunks.append(section(S("Education", "Formação")))
        for edu in data["education"]:
            head = " · ".join(x for x in [edu.get("degree"), edu.get("school")] if x)
            dates = f"<span class='dates'>{esc(edu.get('dates',''))}</span>" if edu.get("dates") else ""
            chunks.append(f"<div class='edu'><span class='head'>{esc(head)}{dates}</span></div>")

    if data.get("certifications"):
        chunks.append(section(S("Certifications", "Certificações")))
        chunks.append(f"<div class='certs'>{esc(', '.join(data['certifications']))}</div>")

    chunks.append("</body></html>")
    html_doc = "".join(chunks)
    weasyprint.HTML(string=html_doc, base_url=".").write_pdf(str(path))
    return str(path)

# ---------------------------------------------------------------------------
# Verificação ATS
# ---------------------------------------------------------------------------
def verify_pdf(pdf_path, name):
    from pypdf import PdfReader
    reader = PdfReader(str(pdf_path))
    pages = len(reader.pages)
    text = "\n".join((p.extract_text() or "") for p in reader.pages)
    ok_text = len(text.strip()) > 50
    ok_name = (name or "").strip() in text if name else True
    print(f"[verify] PDF: {pages} página(s), {len(text)} chars extraídos, "
          f"nome presente: {ok_name}")
    if not ok_text:
        print("[verify] ⚠️  PDF com pouco texto extraível — pode não ser ATS-parseável!")
        return False
    if not ok_name:
        print("[verify] ⚠️  Nome do candidato não encontrado no texto do PDF.")
        return False
    return True

# ---------------------------------------------------------------------------
def main():
    if len(sys.argv) < 3:
        print("Uso: build_resume.py <input.json> <output_base> [--letter]")
        sys.exit(2)
    in_path = Path(sys.argv[1])
    base = sys.argv[2]
    letter = "--letter" in sys.argv[3:]

    data = json.loads(in_path.read_text(encoding="utf-8"))
    name = data.get("meta", {}).get("name", "")

    docx_path = build_docx(data, base + ".docx", letter=letter)
    print(f"[docx] OK: {docx_path}")

    try:
        pdf_path = build_pdf(data, base + ".pdf", letter=letter)
        print(f"[pdf] OK: {pdf_path}")
        ok = verify_pdf(pdf_path, name)
        print("[ats] ✅ VERIFICADO" if ok else "[ats] ❌ FALHOU VERIFICAÇÃO")
    except Exception as e:  # noqa: BLE001
        print(f"[pdf] ❌ FALHOU: {e}")
        print("[pdf] Fallback: tente LibreOffice (`soffice --headless --convert-to pdf <docx>`)")
        ok = False

    sys.exit(0 if ok else 1)

if __name__ == "__main__":
    main()
